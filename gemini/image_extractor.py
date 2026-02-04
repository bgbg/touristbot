"""
Image Extraction Module for DOCX Files

Extracts images from MS Word documents along with their captions and context.
Handles Hebrew text and RTL content correctly.
"""

import io
import os
from dataclasses import dataclass
from typing import List, Optional, Tuple

from docx import Document
from docx.oxml import CT_Blip
from docx.oxml.xmlchemy import BaseOxmlElement
from PIL import Image


@dataclass
class ExtractedImage:
    """Represents an extracted image with its metadata"""

    image_data: bytes
    image_format: str  # 'png', 'jpeg', 'jpg', etc.
    caption: str
    context_before: str
    context_after: str
    paragraph_index: int  # Index in document for reference


class ImageExtractor:
    """Extracts images and metadata from DOCX files"""

    # Image validation limits
    MAX_IMAGE_SIZE_MB = 1
    VALID_IMAGE_FORMATS = {'jpg', 'jpeg', 'png', 'gif', 'webp', 'heic', 'heif'}

    def __init__(self, docx_path: str):
        """
        Initialize image extractor

        Args:
            docx_path: Path to DOCX file
        """
        self.docx_path = docx_path
        self.document = Document(docx_path)

    def extract_images(self) -> List[ExtractedImage]:
        """
        Extract all images from DOCX file with captions and context

        Returns:
            List of ExtractedImage objects

        Raises:
            ValueError: If DOCX file cannot be parsed
        """
        extracted_images = []

        # Iterate through paragraphs to find images
        for i, paragraph in enumerate(self.document.paragraphs):
            # Check if paragraph contains images (via runs)
            for run in paragraph.runs:
                # Look for inline images in the run's XML using findall with full namespace
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')

                for drawing in drawings:
                    pics = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic')

                    for shape in pics:
                        # Extract image data from relationship
                        image_data, image_format = self._extract_image_data(shape)

                        if image_data:
                            # Extract caption (next paragraph after image)
                            caption = self._get_caption(i)

                            # Extract context paragraphs
                            context_before = self._get_paragraph_text(i - 1)
                            # Skip the caption paragraph when getting context after
                            context_after = self._get_paragraph_text(i + 2)

                            extracted_images.append(
                                ExtractedImage(
                                    image_data=image_data,
                                    image_format=image_format,
                                    caption=caption,
                                    context_before=context_before,
                                    context_after=context_after,
                                    paragraph_index=i,
                                )
                            )

        return extracted_images

    def _scale_image(self, image_data: bytes, image_format: str, target_size_mb: float = 1.0) -> bytes:
        """
        Scale down an image to meet the target size limit

        Args:
            image_data: Original image bytes
            image_format: Image format (jpg, png, etc.)
            target_size_mb: Target size in MB (default: 10MB)

        Returns:
            Scaled image bytes
        """
        try:
            # Open image from bytes
            img = Image.open(io.BytesIO(image_data))

            # Convert RGBA to RGB for JPEG format (JPEG doesn't support transparency)
            if image_format in ('jpg', 'jpeg') and img.mode in ('RGBA', 'LA', 'P'):
                # Create white background
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                img = background

            # Start with original dimensions
            width, height = img.size
            quality = 85  # Initial quality for JPEG

            # Iteratively reduce size until it fits
            max_iterations = 10
            for iteration in range(max_iterations):
                output = io.BytesIO()

                # Save with current dimensions and quality
                if image_format in ('jpg', 'jpeg'):
                    img_resized = img.resize((width, height), Image.Resampling.LANCZOS)
                    img_resized.save(output, format='JPEG', quality=quality, optimize=True)
                elif image_format == 'png':
                    img_resized = img.resize((width, height), Image.Resampling.LANCZOS)
                    img_resized.save(output, format='PNG', optimize=True)
                else:
                    # For other formats, try JPEG conversion
                    img_resized = img.resize((width, height), Image.Resampling.LANCZOS)
                    if img_resized.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img_resized.size, (255, 255, 255))
                        if img_resized.mode == 'P':
                            img_resized = img_resized.convert('RGBA')
                        background.paste(img_resized, mask=img_resized.split()[-1] if img_resized.mode in ('RGBA', 'LA') else None)
                        img_resized = background
                    img_resized.save(output, format='JPEG', quality=quality, optimize=True)

                output.seek(0)
                scaled_data = output.read()
                scaled_size_mb = len(scaled_data) / (1024 * 1024)

                # Check if size is acceptable
                if scaled_size_mb <= target_size_mb:
                    print(f"Image scaled successfully: {len(image_data)/(1024*1024):.1f}MB -> {scaled_size_mb:.1f}MB "
                          f"(dimensions: {img.size[0]}x{img.size[1]} -> {width}x{height}, quality: {quality})")
                    return scaled_data

                # Reduce dimensions or quality for next iteration
                if quality > 60:
                    quality -= 10
                else:
                    width = int(width * 0.85)
                    height = int(height * 0.85)
                    quality = 85  # Reset quality when reducing dimensions

                # Safety check - don't go too small
                if width < 200 or height < 200:
                    print(f"Warning: Image scaled to minimum dimensions, size is {scaled_size_mb:.1f}MB")
                    return scaled_data

            # If we couldn't get it small enough after max iterations, return the last attempt
            print(f"Warning: Could not scale image below {target_size_mb}MB after {max_iterations} iterations, "
                  f"returning image at {scaled_size_mb:.1f}MB")
            return scaled_data

        except Exception as e:
            print(f"Warning: Could not scale image: {e}, returning original")
            return image_data

    def _extract_image_data(self, pic_element: BaseOxmlElement) -> Tuple[Optional[bytes], str]:
        """
        Extract image binary data and format from picture element

        Args:
            pic_element: Picture XML element

        Returns:
            Tuple of (image_data, format) or (None, "") if extraction fails
        """
        try:
            # Find the blip element (contains image reference) using findall
            blips = pic_element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')

            if not blips:
                return None, ""

            blip = blips[0]

            # Get the relationship ID
            embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")

            if not embed_id:
                return None, ""

            # Get the image part from the document relationships
            image_part = self.document.part.related_parts[embed_id]

            # Extract image data and format
            image_data = image_part.blob
            image_format = image_part.content_type.split('/')[-1].lower()  # e.g., 'image/jpeg' -> 'jpeg'

            # Normalize format
            if image_format == 'jpeg':
                image_format = 'jpg'

            # Validate format
            if image_format not in self.VALID_IMAGE_FORMATS:
                print(f"Warning: Invalid image format '{image_format}', skipping")
                return None, ""

            # Check size and scale if necessary
            size_mb = len(image_data) / (1024 * 1024)
            if size_mb > self.MAX_IMAGE_SIZE_MB:
                print(f"Image size {size_mb:.1f}MB exceeds {self.MAX_IMAGE_SIZE_MB}MB limit, scaling down...")
                image_data = self._scale_image(image_data, image_format, target_size_mb=self.MAX_IMAGE_SIZE_MB)
                # After scaling, format is normalized to jpg for most cases
                if image_format not in ('jpg', 'png'):
                    image_format = 'jpg'

            return image_data, image_format

        except Exception as e:
            print(f"Warning: Could not extract image data: {e}")
            return None, ""

    def _get_caption(self, image_paragraph_index: int) -> str:
        """
        Get caption for image (paragraph immediately after image)

        Args:
            image_paragraph_index: Index of paragraph containing image

        Returns:
            Caption text or empty string
        """
        caption_index = image_paragraph_index + 1

        if caption_index < len(self.document.paragraphs):
            caption = self.document.paragraphs[caption_index].text.strip()
            return caption

        return ""

    def _get_paragraph_text(self, index: int) -> str:
        """
        Get text of paragraph at index

        Args:
            index: Paragraph index

        Returns:
            Paragraph text or empty string if out of bounds
        """
        if 0 <= index < len(self.document.paragraphs):
            return self.document.paragraphs[index].text.strip()

        return ""

    def save_images(
        self, images: List[ExtractedImage], output_dir: str, base_name: str = "image"
    ) -> List[str]:
        """
        Save extracted images to disk with sequential naming

        Args:
            images: List of ExtractedImage objects
            output_dir: Directory to save images
            base_name: Base name for files (default: "image")

        Returns:
            List of saved file paths

        Raises:
            OSError: If directory creation or file writing fails
        """
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)

        saved_paths = []

        for i, image in enumerate(images, 1):
            # Generate filename: image_001.jpg, image_002.png, etc.
            filename = f"{base_name}_{i:03d}.{image.image_format}"
            filepath = os.path.join(output_dir, filename)

            # Write image data to file
            with open(filepath, 'wb') as f:
                f.write(image.image_data)

            saved_paths.append(filepath)

        return saved_paths


def extract_images_from_docx(
    docx_path: str, output_dir: Optional[str] = None
) -> List[ExtractedImage]:
    """
    Convenience function to extract images from DOCX file

    Args:
        docx_path: Path to DOCX file
        output_dir: Optional directory to save images

    Returns:
        List of ExtractedImage objects

    Example:
        >>> images = extract_images_from_docx('document.docx', 'output/images')
        >>> for img in images:
        ...     print(f"Caption: {img.caption}")
        ...     print(f"Format: {img.image_format}")
    """
    extractor = ImageExtractor(docx_path)
    images = extractor.extract_images()

    if output_dir:
        extractor.save_images(images, output_dir)

    return images
