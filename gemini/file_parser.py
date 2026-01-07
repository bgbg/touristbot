"""
File Parser - Extract text content from various file formats
Supports: .txt, .md, .docx, .pdf
"""

import os
from typing import Optional


def parse_file(file_path: str) -> str:
    """
    Parse file and extract text content

    Args:
        file_path: Path to file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file format is not supported
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return parse_txt(file_path)
    elif ext == ".md":
        return parse_markdown(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def parse_txt(file_path: str) -> str:
    """Parse plain text file"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_markdown(file_path: str) -> str:
    """Parse markdown file"""
    try:
        import markdown

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            md_content = f.read()

        # Convert markdown to plain text (strip HTML tags)
        html = markdown.markdown(md_content)

        # Simple HTML tag removal
        import re

        text = re.sub("<[^<]+?>", "", html)

        return text
    except ImportError:
        # Fallback: treat as plain text
        print(
            f"Warning: markdown library not found, treating {file_path} as plain text"
        )
        return parse_txt(file_path)


def parse_docx(file_path: str) -> str:
    """Parse DOCX file"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    except ImportError:
        raise ValueError(
            f"python-docx library required for .docx files. "
            f"Install with: pip install python-docx"
        )
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file {file_path}: {e}")


def parse_pdf(file_path: str) -> str:
    """Parse PDF file"""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)

        return "\n\n".join(text_parts)

    except ImportError:
        raise ValueError(
            f"PyPDF2 library required for .pdf files. "
            f"Install with: pip install PyPDF2"
        )
    except Exception as e:
        raise ValueError(f"Error parsing PDF file {file_path}: {e}")
